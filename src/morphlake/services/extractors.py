"""Lightweight document extraction and chunking."""

from __future__ import annotations

import io
import json
import re
from contextlib import suppress
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from pypdf import PdfReader

from morphlake.errors import MorphLakeError
from morphlake.models import MediaType

DOCUMENT_EXTENSIONS = {
    ".doc",
    ".docx",
    ".pdf",
    ".txt",
    ".md",
    ".csv",
    ".json",
    ".html",
    ".htm",
}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}
AUDIO_EXTENSIONS = {".wav", ".mp3", ".m4a", ".flac", ".ogg", ".aac"}


@dataclass(frozen=True)
class TextChunk:
    index: int
    start: int
    end: int
    text: str


def classify(filename: str) -> MediaType:
    suffix = Path(filename).suffix.lower()
    if suffix in DOCUMENT_EXTENSIONS:
        return "document"
    if suffix in IMAGE_EXTENSIONS:
        return "image"
    if suffix in AUDIO_EXTENSIONS:
        return "audio"
    raise MorphLakeError("unsupported_media_type", f"Unsupported file extension: {suffix}", 415)


def extract_text(filename: str, body: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        reader = PdfReader(io.BytesIO(body))
        return "\n\n".join(page.extract_text() or "" for page in reader.pages).strip()
    if suffix == ".docx":
        document = Document(io.BytesIO(body))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            paragraphs.extend("\t".join(cell.text for cell in row.cells) for row in table.rows)
        return "\n".join(paragraphs).strip()
    if suffix == ".doc":
        raise MorphLakeError(
            "legacy_word_not_supported",
            "Binary .doc files must be converted to .docx before upload",
            415,
        )
    text = body.decode("utf-8-sig", errors="replace")
    if suffix == ".json":
        with suppress(json.JSONDecodeError):
            text = json.dumps(json.loads(text), ensure_ascii=False, indent=2)
    if suffix in {".html", ".htm"}:
        text = re.sub(r"<script\b[^>]*>.*?</script>", " ", text, flags=re.I | re.S)
        text = re.sub(r"<style\b[^>]*>.*?</style>", " ", text, flags=re.I | re.S)
        text = re.sub(r"<[^>]+>", " ", text)
    return re.sub(r"[ \t]+", " ", text).strip()


def chunk_text(text: str, size: int, overlap: int) -> list[TextChunk]:
    if size <= 0 or overlap < 0 or overlap >= size:
        raise ValueError("chunk size must be positive and overlap must be in [0, size)")
    normalized = text.strip()
    if not normalized:
        return []
    chunks: list[TextChunk] = []
    start = 0
    while start < len(normalized):
        hard_end = min(start + size, len(normalized))
        end = hard_end
        if hard_end < len(normalized):
            candidates = [
                normalized.rfind(separator, start + size // 2, hard_end)
                for separator in ("\n\n", "\n", "。", ". ", " ")
            ]
            best = max(candidates)
            if best > start:
                end = best + 1
        value = normalized[start:end].strip()
        if value:
            chunks.append(TextChunk(len(chunks), start, end, value))
        if end >= len(normalized):
            break
        start = max(end - overlap, start + 1)
    return chunks
